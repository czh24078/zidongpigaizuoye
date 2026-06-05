import asyncio
import json
import logging
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy import select, delete, or_, func
from sqlalchemy.ext.asyncio import AsyncSession

from src.config import config
from src.database import get_db
from src.models.schemas import (
    CorrectionResponse,
    HistoryItem,
    ExamResponse,
    UpdateAnswersRequest,
    CorrectionDetail,
    QuestionBankItemSchema,
    AddToBankRequest,
    AIGenerateRequest,
    AIGenerateResponse,
    AIGeneratedQuestion,
)
from src.models.db_models import (
    Exam,
    Question,
    Correction,
    CorrectionDetail as CorrectionDetailDB,
    QuestionBankItem,
)

logger = logging.getLogger(__name__)

try:
    from src.agents.homework_agent import HomeworkAgent
    homework_agent = HomeworkAgent()
    AGENT_AVAILABLE = True
except Exception:
    AGENT_AVAILABLE = False
    homework_agent = None

router = APIRouter()

# 记录文件目录
BASE_DIR = Path(__file__).resolve().parent.parent.parent
RECORDS_DIR = BASE_DIR / "uploads" / "records"
RECORDS_DIR.mkdir(parents=True, exist_ok=True)


# =====================================================================
# Mock 数据
# =====================================================================

MOCK_CORRECTION_RESULT = """# 作业批改报告

## 总分：85 / 100

> 整体掌握良好，注意符号细节与规范。

## 逐题批改

| 题号 | 学生答案 | 标准答案 | 判定 | 得分 | 分析 |
|------|----------|----------|------|------|------|
| 1 | 3.14 | 3.14 | ✅ | 4/4 | 答案正确。 |
| 2 | 5 | 6 | ❌ | 2/4 | 计算时符号出错，过程正确。 |
| 3 | B | B | ✅ | 4/4 | 概念清晰。 |

> 本次批改由 AI 辅助完成，如有疑问请咨询任课老师。
"""

MOCK_EXAM_QUESTIONS = [
    {"question_no": "1", "question_text": "π 保留两位小数等于？", "standard_answer": "3.14", "analysis": "直接取近似值。"},
    {"question_no": "2", "question_text": "2 + 2 × 2 = ?", "standard_answer": "6", "analysis": "先乘后加。"},
    {"question_no": "3", "question_text": "下列哪一个是平行四边形的判定条件？", "standard_answer": "B", "analysis": "对边平行且相等。"},
]

MOCK_CORRECTION_DETAILS = [
    {"question_no": "1", "question_text": "π 保留两位小数等于？", "student_answer": "3.14", "standard_answer": "3.14",
     "is_correct": True, "score": 4, "full_score": 4, "analysis": "答案正确。"},
    {"question_no": "2", "question_text": "2 + 2 × 2 = ?", "student_answer": "5", "standard_answer": "6",
     "is_correct": False, "score": 2, "full_score": 4, "analysis": "先乘后加的运算顺序未正确应用，过程部分分。"},
    {"question_no": "3", "question_text": "判定平行四边形", "student_answer": "B", "standard_answer": "B",
     "is_correct": True, "score": 4, "full_score": 4, "analysis": "概念掌握到位。"},
]


# =====================================================================
# 工具
# =====================================================================

def _validate_image_file(file: UploadFile) -> None:
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    ext = os.path.splitext(file.filename)[1].lower().lstrip(".")
    if ext == "jpeg":
        ext = "jpg"
    if ext not in config.ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型：.{ext}，仅支持 jpg/png/webp")


async def _save_upload_file(file: UploadFile, subdir: str = "") -> str:
    target_dir = os.path.join(config.UPLOAD_DIR, subdir) if subdir else config.UPLOAD_DIR
    os.makedirs(target_dir, exist_ok=True)
    file_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename or "")[1].lower() or ".jpg"
    save_path = os.path.join(target_dir, f"{file_id}{ext}")
    content = await file.read()
    if len(content) > config.MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="文件大小超过 10MB 限制")
    with open(save_path, "wb") as f:
        f.write(content)
    return save_path


async def _auto_add_to_bank(session: AsyncSession, exam_id: Optional[str], questions: list[dict], display_name: str) -> int:
    """将题目自动加入题库（去重），返回新增数量。"""
    max_no_result = await session.execute(select(func.max(QuestionBankItem.bank_no)))
    max_no = max_no_result.scalar() or 0
    added = 0
    for q in questions:
        q_text = q.get("question_text", "")
        q_answer = q.get("standard_answer", "")
        existing = await session.execute(
            select(QuestionBankItem).where(
                QuestionBankItem.question_text == q_text,
                QuestionBankItem.standard_answer == q_answer,
            )
        )
        if existing.scalar_one_or_none():
            continue
        subject = q.get("subject") or _guess_subject(display_name, q_text)
        added += 1
        session.add(QuestionBankItem(
            exam_id=exam_id,
            question_no=q.get("question_no", ""),
            question_text=q_text,
            standard_answer=q_answer,
            analysis=q.get("analysis", ""),
            exam_filename=display_name,
            bank_no=max_no + added,
            subject=subject,
            added_at=datetime.now(),
        ))
    return added


_SUBJECT_KEYWORDS = {
    "语文": [
        # 字词基础
        "拼音", "汉字", "笔画", "偏旁", "部首", "查字典", "音序", "音节",
        "多音字", "形近字", "同音字", "多义字", "错别字", "生字", "认读",
        "成语", "谚语", "歇后语", "格言", "名言", "警句", "对联", "春联",
        "近义词", "反义词", "同义词", "褒义词", "贬义词", "中性词",
        "词语搭配", "词语解释", "关联词", "连词", "量词", "代词",
        # 句子
        "造句", "组词", "扩句", "缩句", "句式", "陈述句", "疑问句", "祈使句", "感叹句",
        "把字句", "被字句", "反问句", "设问句", "双重否定句", "转述句",
        "修改病句", "病句", "语病", "搭配不当", "重复啰嗦",
        "选词填空", "关联词填空", "补全句子", "连词成句", "排序",
        # 修辞
        "修辞", "比喻", "拟人", "排比", "夸张", "对偶", "设问", "反问",
        "借代", "对比", "反复", "引用", "通感", "象征", "衬托",
        # 标点
        "标点", "逗号", "句号", "问号", "感叹号", "冒号", "分号", "引号",
        # 古诗文
        "古诗", "唐诗", "宋词", "元曲", "诗经", "楚辞", "汉乐府", "古文",
        "文言文", "背诵", "默写", "填空", "名句", "千古名句",
        "李白", "杜甫", "白居易", "王维", "孟浩然", "苏轼", "辛弃疾", "陆游",
        "陶渊明", "王之涣", "王昌龄", "杜牧", "李商隐", "柳宗元", "韩愈",
        "范仲淹", "欧阳修", "王安石", "李清照", "文天祥", "屈原",
        "诗人", "词人", "作者", "朝代", "诗题", "题目", "题材",
        "五言", "七言", "绝句", "律诗", "小令", "长调",
        # 文学
        "《", "》", "论语", "孟子", "大学", "中庸", "庄子", "老子", "韩非子",
        "红楼梦", "三国演义", "水浒传", "西游记", "聊斋志异", "儒林外史",
        "鲁迅", "老舍", "巴金", "冰心", "朱自清", "矛盾", "郭沫若", "沈从文",
        "小说", "散文", "戏剧", "诗歌", "寓言", "童话", "神话", "传说",
        "记叙文", "说明文", "议论文", "应用文", "新闻", "通讯", "报告文学",
        # 阅读与写作
        "阅读理解", "阅读", "短文", "选文", "文章", "课文", "全文", "段落",
        "中心思想", "主要内容", "概括", "归纳", "总结", "分析", "评价",
        "人物形象", "性格特点", "写作手法", "表达方式", "描写方法",
        "环境描写", "心理描写", "外貌描写", "动作描写", "神态描写", "语言描写",
        "正面描写", "侧面描写", "细节描写", "白描", "渲染", "烘托",
        "顺叙", "倒叙", "插叙", "补叙", "伏笔", "照应", "悬念",
        "开门见山", "首尾呼应", "承上启下", "画龙点睛", "卒章显志",
        "作文", "写作", "日记", "周记", "书信", "通知", "启事",
        "请假条", "留言条", "读后感", "观后感", "演讲稿", "倡议书",
        "审题", "立意", "选材", "提纲", "开头", "结尾", "过渡",
        # 综合
        "看拼音", "写词语", "按要求", "连线", "标序号", "下列", "选出",
        "正确的是", "错误的是", "正确的一项", "不正确的一项",
        "把下列", "改为", "仿写", "续写", "缩写", "扩写", "改写",
        "表达", "思想感情", "体会", "感受", "启发", "道理", "启示",
        "解释", "意思", "含义", "理解", "领略", "欣赏", "品味",
        "判断对错", "判断下列说法", "补充完整", "按课文", "按原文",
    ],
    "数学": [
        # 数与运算
        "计算", "口算", "笔算", "估算", "竖式", "递等式", "脱式", "简便", "简便运算",
        "加法", "减法", "乘法", "除法", "混合运算", "验算", "列式", "算式",
        "整数", "小数", "分数", "百分数", "正数", "负数", "自然数", "奇数", "偶数",
        "质数", "合数", "因数", "倍数", "公因数", "公倍数", "最大公因数", "最小公倍数",
        "约分", "通分", "最简分数", "带分数", "假分数", "真分数",
        "倒数", "相反数", "绝对值", "四舍五入", "近似值", "约等于",
        "等于", "大于", "小于", "比大小", "多少", "一共", "几", "结果",
        "加起来", "减去", "乘以", "除以", "余数", "相差", "增加", "减少",
        "平均", "每", "倍", "份", "除以", "除", "被除", "除数", "商",
        "四则运算", "运算顺序", "先乘除", "后加减", "括号", "中括号",
        # 数量关系与方程
        "方程", "解方程", "等式", "不等式", "未知数", "x=", "y=", "z=",
        "一元一次", "二元一次", "一元二次", "二次函数", "一次函数",
        "比例", "正比例", "反比例", "比值", "化简比", "求比值",
        # 几何
        "几何", "图形", "平面", "立体", "周长", "面积", "体积", "表面积", "底面积",
        "角度", "边长", "半径", "直径", "高", "底", "斜边", "直角", "锐角", "钝角",
        "三角形", "四边形", "长方形", "正方形", "平行四边形", "梯形", "菱形",
        "五边形", "六边形", "多边形", "圆", "扇形", "椭圆", "环形",
        "正方体", "长方体", "圆柱", "圆锥", "球", "棱柱", "棱锥",
        "对称", "轴对称", "平移", "旋转", "相似", "全等", "位似",
        "勾股定理", "三角函数", "sin", "cos", "tan", "正弦", "余弦", "正切",
        "象限", "坐标", "抛物线", "双曲线", "坐标系", "数轴",
        # 统计与概率
        "统计", "统计表", "统计图", "条形图", "折线图", "扇形图", "直方图",
        "平均数", "中位数", "众数", "方差", "标准差", "极差",
        "概率", "可能性", "随机", "必然", "不可能", "可能",
        # 单位与测量
        "厘米", "毫米", "分米", "米", "千米", "公里", "cm", "mm", "dm", "m", "km",
        "平方米", "平方厘米", "平方千米", "公顷", "m²", "cm²",
        "克", "千克", "公斤", "吨", "g", "kg",
        "毫升", "升", "mL", "L",
        "元", "角", "分", "人民币",
        "时", "分", "秒", "小时", "分钟", "分钟", "秒钟",
        "换算", "单位换算", "进率",
        # 应用题
        "应用题", "解决问题", "列式解答", "用方程解", "用比例解",
        "路程", "速度", "时间", "相遇", "追及", "行船", "过桥",
        "工程", "工作效率", "工作时间", "工作总量",
        "单价", "数量", "总价", "打折", "利润", "利率", "利息",
        "植树", "间隔", "方阵", "鸡兔同笼", "盈亏",
        "抽屉", "逻辑推理", "找规律", "周期", "枚举", "排列组合",
        "行程", "浓度", "年龄", "钟表", "余数",
    ],
    "物理": [
        # 力学
        "力", "重力", "弹力", "摩擦力", "浮力", "压力", "支持力", "拉力", "推力",
        "作用力", "反作用力", "平衡力", "合力", "分力", "力的合成", "力的分解",
        "牛顿", "牛顿第一定律", "惯性", "牛顿第二定律", "牛顿第三定律",
        "质量", "重量", "密度", "体积", "kg/m³", "g/cm³",
        "速度", "加速度", "匀速", "变速", "匀速直线", "自由落体",
        "位移", "路程", "时间", "平均速度", "瞬时速度", "m/s", "km/h",
        "功", "功率", "机械功", "有用功", "总功", "额外功", "机械效率",
        "动能", "势能", "重力势能", "弹性势能", "机械能", "能量守恒",
        "杠杆", "滑轮", "定滑轮", "动滑轮", "滑轮组", "斜面", "轮轴",
        "压强", "压力", "帕斯卡", "液体压强", "大气压强", "流体压强",
        "浮力", "阿基米德", "浮沉", "漂浮", "悬浮", "下沉", "上浮",
        "天平", "弹簧", "弹簧测力计", "量筒", "刻度尺", "秒表",
        # 光学
        "光", "光的传播", "直线传播", "光速", "光源", "光线",
        "反射", "反射定律", "镜面反射", "漫反射", "入射角", "反射角",
        "折射", "折射定律", "折射角", "全反射", "临界角",
        "透镜", "凸透镜", "凹透镜", "焦距", "焦点", "物距", "像距",
        "实像", "虚像", "放大镜", "显微镜", "望远镜", "照相机", "投影仪",
        "光的色散", "光谱", "红外线", "紫外线", "可见光",
        # 电学
        "电路", "串联", "并联", "混联", "通路", "断路", "短路",
        "电流", "电压", "电阻", "欧姆", "安培", "伏特", "欧姆定律",
        "电源", "用电器", "开关", "导线", "电池", "干电池", "蓄电池",
        "正极", "负极", "正电荷", "负电荷", "电子", "自由电子",
        "导体", "绝缘体", "半导体", "超导体",
        "电功", "电功率", "电能", "千瓦时", "焦耳", "瓦特",
        "电流表", "电压表", "滑动变阻器", "万用表", "验电器",
        "电磁", "磁场", "磁感线", "磁极", "电磁铁", "电磁感应",
        "安培定则", "左手定则", "右手定则", "楞次定律",
        # 热学
        "温度", "热量", "比热容", "热值", "内能", "热传递", "热平衡",
        "熔化", "凝固", "汽化", "液化", "升华", "凝华", "蒸发", "沸腾",
        "熔点", "沸点", "凝固点", "晶体", "非晶体",
        "分子", "原子", "扩散", "热运动", "分子间作用力",
        "温度计", "体温计", "酒精温度计",
        # 声学
        "声音", "声波", "音调", "响度", "音色", "频率", "振幅",
        "超声波", "次声波", "噪声", "乐音", "分贝",
        "声速", "回声", "共鸣", "振动", "介质",
        # 综合
        "物理量", "单位", "测量", "实验", "误差", "有效数字",
    ],
    "历史": [
        # 中国古代
        "朝代", "皇帝", "帝王", "天子", "君主", "统治", "王朝", "帝国", "政权",
        "秦朝", "汉朝", "唐朝", "宋朝", "元朝", "明朝", "清朝",
        "隋朝", "三国", "魏晋", "南北朝", "五代十国", "西夏", "辽", "金",
        "秦始皇", "汉武帝", "唐太宗", "宋太祖", "成吉思汗", "忽必烈",
        "朱元璋", "康熙", "乾隆", "雍正", "光绪", "溥仪",
        "统一", "分裂", "割据", "兼并", "扩张", "版图", "疆域",
        "制度", "封建", "中央集权", "君主专制", "郡县制", "分封制",
        "丞相", "宰相", "尚书", "御史", "刺史", "太守", "县令",
        "科举", "察举", "九品中正", "八股文", "进士", "状元",
        "变法", "改革", "商鞅", "王安石", "张居正", "百日维新",
        "焚书坑儒", "罢黜百家", "独尊儒术", "推恩令",
        "三省六部", "行省制度", "内阁", "军机处", "文字狱",
        "丝绸之路", "郑和", "下西洋", "朝贡", "贸易",
        "长城", "大运河", "都江堰", "灵渠", "故宫", "兵马俑",
        # 思想文化
        "儒家", "道家", "法家", "墨家", "兵家", "纵横家", "阴阳",
        "孔子", "孟子", "荀子", "庄子", "老子", "墨子", "韩非子",
        "董仲舒", "朱熹", "王阳明", "程朱理学", "陆王心学",
        # 重要事件
        "陈胜吴广", "大泽乡", "楚汉之争", "文景之治", "光武中兴",
        "贞观之治", "开元盛世", "安史之乱", "黄巢起义",
        "陈桥兵变", "杯酒释兵权", "澶渊之盟", "靖康之耻", "岳飞",
        "赤壁之战", "淝水之战", "官渡之战", "巨鹿之战",
        "靖难之役", "土木堡之变", "郑成功", "收复台湾",
        # 近现代
        "鸦片战争", "虎门销烟", "林则徐", "南京条约", "不平等条约",
        "太平天国", "洪秀全", "洋务运动", "李鸿章", "曾国藩", "左宗棠",
        "甲午", "甲午战争", "马关条约", "义和团", "八国联军", "辛丑条约",
        "辛亥革命", "孙中山", "三民主义", "中华民国", "袁世凯", "北洋军阀",
        "新文化运动", "五四运动", "五四", "陈独秀", "胡适",
        "共产党", "共产主义", "南昌起义", "秋收起义", "井冈山", "长征", "遵义会议",
        "抗日战争", "全面抗战", "七七事变", "九一八", "南京大屠杀",
        "解放战争", "三大战役", "渡江战役", "新中国成立", "开国大典",
        "毛泽东", "周恩来", "朱德", "刘少奇", "邓小平", "彭德怀",
        # 世界历史
        "古埃及", "古巴比伦", "古印度", "古希腊", "古罗马", "雅典", "斯巴达",
        "文艺复兴", "启蒙运动", "宗教改革", "地理大发现", "哥伦布", "麦哲伦",
        "工业革命", "蒸汽机", "瓦特", "资产阶级革命",
        "美国独立", "独立宣言", "华盛顿", "法国大革命", "拿破仑",
        "南北战争", "林肯", "明治维新", "俄国革命", "十月革命",
        "世界大战", "第一次世界", "第二次世界", "冷战", "联合国",
        # 综合概念
        "战争", "革命", "起义", "政变", "改革", "变法", "维新",
        "条约", "协定", "宣言", "公告", "媾和", "停战",
        "封建", "殖民", "独立", "解放", "建国", "成立", "灭亡",
        "民族", "国家", "领土", "主权", "制度", "体制", "政体",
        "民主", "共和", "君主", "专制", "独裁", "立宪",
        "宪法", "议会", "选举", "国会", "内阁", "总统", "首相",
        "公元前", "公元", "世纪", "年代", "时期", "年间", "初年", "末年",
        "导火线", "根本原因", "直接原因", "主要影响", "历史意义",
    ],
}


def _guess_subject(filename: str = "", question_text: str = "") -> str:
    """根据题目文本和文件名综合判断科目，无法判断时返回'其他'。"""
    text = (question_text + " " + filename).lower()
    scores = {}
    for subj, keywords in _SUBJECT_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw.lower() in text)
        if score > 0:
            scores[subj] = score

    if not scores:
        return "其他"
    best = max(scores, key=scores.get)
    return best if scores[best] >= 2 else "其他"


def _is_valid_question(question_text: str, standard_answer: str) -> bool:
    """判断题目是否完整有效，过滤残缺题目。"""
    if not question_text or not standard_answer:
        return False
    qt = question_text.strip()
    sa = standard_answer.strip()
    # 题干或答案太短（<5字符）视为不完整
    if len(qt) < 5 or len(sa) < 1:
        return False
    # 题干仅含单个字符/数字/符号视为不完整
    if len(qt) < 3 and not any('一' <= c <= '鿿' for c in qt):
        return False
    return True


def _save_record_docx(
    correction_id: str, filename: str, exam: Optional[ExamResponse],
    score: Optional[int], details: list[CorrectionDetail],
    result_markdown: str, created_at: datetime,
) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("作业批改记录", level=1)

    info = [
        ("批改ID", correction_id),
        ("作业文件", filename),
        ("批改时间", created_at.strftime("%Y-%m-%d %H:%M:%S")),
    ]
    if exam is not None:
        info.append(("关联试题", f"{exam.id}（{exam.filename}）"))
    if score is not None:
        info.append(("总分", f"{score} / 100"))

    for label, value in info:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(value)

    doc.add_heading("逐题批改记录", level=2)

    if details:
        for i, d in enumerate(details, start=1):
            doc.add_heading(f"第 {d.question_no or i} 题", level=3)
            if d.question_text:
                p = doc.add_paragraph()
                p.add_run("题干：").bold = True
                p.add_run(d.question_text)
            p = doc.add_paragraph()
            p.add_run("学生答案：").bold = True
            p.add_run(d.student_answer or "（未作答/无法识别）")
            p = doc.add_paragraph()
            p.add_run("正确答案：").bold = True
            p.add_run(d.standard_answer)
            judge = "正确" if d.is_correct else "错误"
            if d.score is not None and d.full_score is not None:
                judge += f"（{d.score}/{d.full_score}）"
            p = doc.add_paragraph()
            p.add_run("判定：").bold = True
            p.add_run(judge)
            p = doc.add_paragraph()
            p.add_run("批改分析：").bold = True
            p.add_run(d.analysis)
    else:
        doc.add_paragraph("本次批改未产生结构化的逐题数据，下面是原始批改输出：")
        doc.add_paragraph(result_markdown)

    doc.add_paragraph("—" * 30)
    doc.add_heading("完整批改报告", level=2)
    doc.add_paragraph(result_markdown)

    path = RECORDS_DIR / f"{correction_id}.docx"
    doc.save(str(path))
    return str(path.relative_to(BASE_DIR)).replace("\\", "/")


def _details_from_dicts(items: list[dict]) -> list[CorrectionDetail]:
    out = []
    for d in items or []:
        try:
            out.append(CorrectionDetail(
                question_no=str(d.get("question_no", "")),
                question_text=str(d.get("question_text", "")),
                student_answer=str(d.get("student_answer", "")),
                standard_answer=str(d.get("standard_answer", "")),
                is_correct=bool(d.get("is_correct", False)),
                score=d.get("score"),
                full_score=d.get("full_score"),
                analysis=str(d.get("analysis", "")),
            ))
        except Exception:
            continue
    return out


async def _get_exam(session: AsyncSession, exam_id: str) -> Optional[Exam]:
    result = await session.execute(select(Exam).where(Exam.id == exam_id))
    return result.scalar_one_or_none()


async def _get_exam_pydantic(session: AsyncSession, exam_id: str) -> Optional[ExamResponse]:
    db_exam = await _get_exam(session, exam_id)
    if db_exam is None:
        return None
    await session.refresh(db_exam, ["questions"])
    return db_exam.to_pydantic()


# =====================================================================
# 试题（Exam）接口
# =====================================================================

@router.post("/exam/upload")
async def upload_exam(files: list[UploadFile] = File(...), session: AsyncSession = Depends(get_db)):
    if not files:
        raise HTTPException(status_code=400, detail="请至少上传一张图片")

    saved_paths: list[str] = []
    filenames: list[str] = []
    for file in files:
        _validate_image_file(file)
        path = await _save_upload_file(file, subdir="exams")
        saved_paths.append(path)
        filenames.append(file.filename or "exam.jpg")

    questions_data: list[dict] = []
    if AGENT_AVAILABLE and homework_agent is not None:
        try:
            questions_data = await homework_agent.extract_exam(saved_paths)
        except Exception:
            questions_data = MOCK_EXAM_QUESTIONS
    else:
        questions_data = MOCK_EXAM_QUESTIONS

    if not questions_data:
        raise HTTPException(status_code=500, detail="试题识别失败，请更换清晰的图片重试")

    display_name = " + ".join(filenames) if len(filenames) > 1 else filenames[0]
    exam_id = str(uuid.uuid4())

    # 创建试题记录 + 同时加入题库
    db_exam = Exam(
        id=exam_id, filename=display_name, source="ai", created_at=datetime.now(),
        questions=[Question(
            question_no=q.get("question_no", str(i + 1)),
            question_text=q.get("question_text", ""),
            standard_answer=q.get("standard_answer", ""),
            analysis=q.get("analysis", ""),
            subject=q.get("subject") or _guess_subject(question_text=q.get("question_text", "")),
        ) for i, q in enumerate(questions_data)],
    )
    session.add(db_exam)
    bank_added = await _auto_add_to_bank(session, exam_id, questions_data, display_name)

    await session.commit()
    return {
        "message": f"试题识别完成，{bank_added} 道题目已导入题库",
        "bank_added": bank_added,
        "filename": display_name,
        "exam_id": exam_id,
        "exam": db_exam.to_pydantic().model_dump(mode="json"),
    }


@router.get("/exams")
async def list_exams(
    keyword: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    session: AsyncSession = Depends(get_db),
):
    stmt = select(Exam).order_by(Exam.created_at.desc())

    if start_date:
        try:
            stmt = stmt.where(Exam.created_at >= datetime.fromisoformat(start_date))
        except ValueError:
            raise HTTPException(status_code=400, detail="start_date 格式无效，请使用 ISO 格式（如 2025-01-01）")
    if end_date:
        try:
            stmt = stmt.where(Exam.created_at <= datetime.fromisoformat(end_date))
        except ValueError:
            raise HTTPException(status_code=400, detail="end_date 格式无效，请使用 ISO 格式（如 2025-12-31）")

    result = await session.execute(stmt)
    exams = result.scalars().all()
    exam_list = [e.to_pydantic() for e in exams]

    if keyword:
        kw = keyword.strip().lower()
        exam_list = [e for e in exam_list if any(
            kw in q.question_text.lower() or kw in q.standard_answer.lower()
            for q in e.questions
        )]

    total = len(exam_list)
    start = (page - 1) * page_size
    return {"items": exam_list[start:start + page_size], "total": total}


@router.get("/exam/{exam_id}", response_model=ExamResponse)
async def get_exam(exam_id: str, session: AsyncSession = Depends(get_db)):
    exam = await _get_exam_pydantic(session, exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="试题不存在")
    return exam


@router.put("/exam/{exam_id}/answers", response_model=ExamResponse)
async def update_exam_answers(exam_id: str, payload: UpdateAnswersRequest, session: AsyncSession = Depends(get_db)):
    db_exam = await _get_exam(session, exam_id)
    if not db_exam:
        raise HTTPException(status_code=404, detail="试题不存在")

    # 删除旧题目，插入新题目
    existing_questions = list(db_exam.questions)
    for q in existing_questions:
        await session.delete(q)
    await session.flush()

    for qi in payload.questions:
        q = Question.from_pydantic(qi)
        q.exam_id = db_exam.id
        session.add(q)

    db_exam.source = "manual"
    await session.commit()
    await session.refresh(db_exam, ["questions"])
    return db_exam.to_pydantic()


@router.delete("/exam/{exam_id}")
async def delete_exam(exam_id: str, session: AsyncSession = Depends(get_db)):
    db_exam = await _get_exam(session, exam_id)
    if db_exam:
        await session.delete(db_exam)
        await session.commit()
    return {"status": "ok"}


# =====================================================================
# 批改接口
# =====================================================================

@router.post("/correct", response_model=CorrectionResponse)
async def correct_homework(
    file: UploadFile = File(...),
    exam_id: Optional[str] = Form(None),
    session: AsyncSession = Depends(get_db),
):
    _validate_image_file(file)
    file_path = await _save_upload_file(file)

    exam: Optional[ExamResponse] = None
    standard_answers = None
    if exam_id:
        exam = await _get_exam_pydantic(session, exam_id)
        if not exam:
            raise HTTPException(status_code=404, detail="指定的试题不存在")
        standard_answers = [q.model_dump() for q in exam.questions]

    result_text = ""
    score: Optional[int] = None
    details_objs: list[CorrectionDetail] = []

    if AGENT_AVAILABLE and homework_agent is not None:
        try:
            result_text, score, raw_details = await homework_agent.correct(
                file_path, standard_answers=standard_answers)
            details_objs = _details_from_dicts(raw_details or [])
        except Exception:
            result_text = MOCK_CORRECTION_RESULT
            if standard_answers:
                details_objs = _details_from_dicts(MOCK_CORRECTION_DETAILS)
    else:
        result_text = MOCK_CORRECTION_RESULT
        if standard_answers:
            details_objs = _details_from_dicts(MOCK_CORRECTION_DETAILS)

    # 过滤掉不完整的题目
    if details_objs:
        before = len(details_objs)
        details_objs = [d for d in details_objs if _is_valid_question(d.question_text, d.standard_answer)]
        if len(details_objs) < before:
            logger.info(f"过滤掉 {before - len(details_objs)} 道不完整题目")

    correction_id = str(uuid.uuid4())
    created_at = datetime.now()

    # 已选试题时直接关联原始试题，不创建新的衍生 Exam
    if exam:
        final_exam_id = exam.id
    elif details_objs:
        # 无试题时，自动创建历史试题记录保存题目
        final_exam_id = str(uuid.uuid4())
        db_history_exam = Exam(
            id=final_exam_id,
            filename=f"[批改] {file.filename or 'unknown'}",
            source="correction",
            created_at=created_at,
            questions=[Question(
                question_no=d.question_no,
                question_text=d.question_text,
                standard_answer=d.standard_answer,
                analysis=d.analysis,
            ) for d in details_objs],
        )
        session.add(db_history_exam)
    else:
        final_exam_id = None

    record_path = _save_record_docx(
        correction_id=correction_id, filename=file.filename or "unknown",
        exam=exam, score=score, details=details_objs,
        result_markdown=result_text, created_at=created_at,
    )

    summary = f"作业批改完成，总分 {score} 分。" if score is not None else "作业批改完成。"
    summary += f"（基于试题 {exam.filename}）" if exam else ""

    db_correction = Correction(
        id=correction_id, filename=file.filename or "unknown",
        result=result_text, score=score, summary=summary,
        exam_id=final_exam_id,
        record_path=record_path, created_at=created_at,
    )
    for d in details_objs:
        db_correction.details.append(CorrectionDetailDB.from_pydantic(d))
    session.add(db_correction)

    # 自动将批改题目加入题库（去重+科目分类）
    if details_objs:
        questions_data = [{
            "question_no": d.question_no,
            "question_text": d.question_text,
            "standard_answer": d.standard_answer,
            "analysis": d.analysis,
            "subject": _guess_subject(question_text=d.question_text),
        } for d in details_objs]
        await _auto_add_to_bank(session, final_exam_id, questions_data,
                                f"[批改] {file.filename or 'unknown'}")

    await session.commit()

    return CorrectionResponse(
        id=correction_id, filename=file.filename or "unknown",
        result=result_text, score=score,
        exam_id=final_exam_id,
        details=details_objs or None,
        record_path=record_path, created_at=created_at,
    )


@router.post("/correct/stream")
async def correct_homework_stream(
    file: UploadFile = File(...),
    exam_id: Optional[str] = Form(None),
    session: AsyncSession = Depends(get_db),
):
    _validate_image_file(file)
    file_path = await _save_upload_file(file)

    exam: Optional[ExamResponse] = None
    standard_answers = None
    if exam_id:
        exam = await _get_exam_pydantic(session, exam_id)
        if not exam:
            raise HTTPException(status_code=404, detail="指定的试题不存在")
        standard_answers = [q.model_dump() for q in exam.questions]

    async def event_generator():
        nonlocal exam
        try:
            yield f"data: {json.dumps({'event': 'start', 'message': '开始批改...'}, ensure_ascii=False)}\n\n"
            full_content = ""
            stream_score: Optional[int] = None
            stream_details: list[dict] = []

            if AGENT_AVAILABLE and homework_agent is not None:
                try:
                    async for chunk in homework_agent.correct_stream(file_path, standard_answers=standard_answers):
                        yield f"data: {chunk}\n\n"
                        try:
                            chunk_data = json.loads(chunk)
                            if chunk_data.get('event') == 'content':
                                full_content += chunk_data.get('text', '')
                            elif chunk_data.get('event') == 'final_text':
                                full_content = chunk_data.get('text', full_content)
                                stream_score = chunk_data.get('score')
                                stream_details = chunk_data.get('details') or []
                        except Exception:
                            pass
                except Exception as e:
                    logger.error(f"流式批改失败: {e}")
                    yield f"data: {json.dumps({'event': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
            else:
                for chunk in _mock_stream_chunks():
                    yield f"data: {chunk}\n\n"
                    await asyncio.sleep(0.1)
                    try:
                        cd = json.loads(chunk)
                        if cd.get("event") == "content":
                            full_content += cd.get("text", "")
                    except Exception:
                        pass

            score = stream_score if AGENT_AVAILABLE else 85
            details_objs = _details_from_dicts(stream_details) if AGENT_AVAILABLE else []

            # 过滤掉不完整的题目
            if details_objs:
                before = len(details_objs)
                details_objs = [d for d in details_objs if _is_valid_question(d.question_text, d.standard_answer)]
                if len(details_objs) < before:
                    logger.info(f"流式批改过滤掉 {before - len(details_objs)} 道不完整题目")

            correction_id = str(uuid.uuid4())
            created_at = datetime.now()
            result_text = full_content

            record_path = _save_record_docx(
                correction_id=correction_id, filename=file.filename or "unknown",
                exam=exam, score=score, details=details_objs,
                result_markdown=result_text, created_at=created_at,
            )

            summary = f"作业批改完成，总分 {score} 分。" if score is not None else "作业批改完成。"
            summary += f"（基于试题 {exam.filename}）" if exam else ""

            from src.database import AsyncSessionLocal
            async with AsyncSessionLocal() as save_session:
                # 已选试题时直接关联原始试题
                if exam:
                    final_exam_id = exam.id
                elif details_objs:
                    from src.models.db_models import Question as QDB
                    final_exam_id = str(uuid.uuid4())
                    db_history_exam = Exam(
                        id=final_exam_id,
                        filename=f"[批改] {file.filename or 'unknown'}",
                        source="correction",
                        created_at=created_at,
                        questions=[QDB(
                            question_no=d.question_no,
                            question_text=d.question_text,
                            standard_answer=d.standard_answer,
                            analysis=d.analysis,
                        ) for d in details_objs],
                    )
                    save_session.add(db_history_exam)
                else:
                    final_exam_id = None

                db_correction = Correction(
                    id=correction_id, filename=file.filename or "unknown",
                    result=result_text, score=score, summary=summary,
                    exam_id=final_exam_id,
                    record_path=record_path, created_at=created_at,
                )
                for d in details_objs:
                    db_correction.details.append(CorrectionDetailDB.from_pydantic(d))
                save_session.add(db_correction)

                # 自动将批改题目加入题库
                if details_objs:
                    questions_data = [{
                        "question_no": d.question_no,
                        "question_text": d.question_text,
                        "standard_answer": d.standard_answer,
                        "analysis": d.analysis,
                        "subject": _guess_subject(question_text=d.question_text),
                    } for d in details_objs]
                    await _auto_add_to_bank(save_session, final_exam_id, questions_data,
                                            f"[批改] {file.filename or 'unknown'}")

                await save_session.commit()

            yield f"data: {json.dumps({'event': 'end', 'message': '批改完成'}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"流式生成器异常: {e}")
            yield f"data: {json.dumps({'event': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


def _mock_stream_chunks():
    """将 Mock 批改报告拆分为 content 事件流，与非流式 MOCK_CORRECTION_RESULT 一致。"""
    lines = MOCK_CORRECTION_RESULT.strip().split("\n")
    chunks = []
    for line in lines:
        chunks.append(json.dumps({"event": "content", "text": line + "\n"}, ensure_ascii=False))
    return chunks


# =====================================================================
# 历史 & 记录文件下载
# =====================================================================

@router.get("/history")
async def get_history(
    page: int = 1,
    page_size: int = 20,
    session: AsyncSession = Depends(get_db),
):
    # 先查总数
    count_result = await session.execute(select(func.count(Correction.id)))
    total = count_result.scalar() or 0
    # 分页查询
    result = await session.execute(
        select(Correction).order_by(Correction.created_at.desc())
        .offset((page - 1) * page_size).limit(page_size)
    )
    corrections = result.scalars().all()
    return {"items": [c.to_history_item() for c in corrections], "total": total}


@router.get("/correction/{correction_id}/record")
async def download_record(correction_id: str):
    path = RECORDS_DIR / f"{correction_id}.docx"
    if not path.exists():
        raise HTTPException(status_code=404, detail="批改记录不存在")
    return FileResponse(
        path, filename=f"correction_{correction_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# =====================================================================
# 试卷导出
# =====================================================================

@router.post("/exam-paper/export")
async def export_exam_paper(payload: list[dict]):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("练习试卷", level=1)
    p = doc.add_paragraph()
    p.add_run(f"共 {len(payload)} 题").bold = True

    # 题目部分
    for i, q in enumerate(payload, start=1):
        doc.add_heading(f"第 {i} 题", level=3)
        doc.add_paragraph(q.get("question_text", ""))

    # 答案部分
    doc.add_paragraph("—" * 30)
    doc.add_heading("参考答案", level=2)
    for i, q in enumerate(payload, start=1):
        doc.add_heading(f"第 {i} 题", level=3)
        p = doc.add_paragraph()
        p.add_run("答案：").bold = True
        p.add_run(q.get("standard_answer", ""))
        analysis = q.get("analysis", "")
        if analysis:
            p2 = doc.add_paragraph()
            p2.add_run("解析：").bold = True
            p2.add_run(analysis)

    path = RECORDS_DIR / f"exam_paper_{uuid.uuid4().hex[:8]}.docx"
    doc.save(str(path))
    return FileResponse(
        path, filename=f"练习试卷_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# =====================================================================
# 题库接口
# =====================================================================

@router.get("/question-bank")
async def list_question_bank(
    keyword: Optional[str] = None,
    question_no: Optional[str] = None,
    subject: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    session: AsyncSession = Depends(get_db),
):
    stmt = select(QuestionBankItem).order_by(QuestionBankItem.bank_no.asc())
    count_stmt = select(func.count(QuestionBankItem.id))

    if keyword:
        kw = keyword.strip()
        cond = or_(
            QuestionBankItem.question_text.contains(kw),
            QuestionBankItem.standard_answer.contains(kw),
            QuestionBankItem.analysis.contains(kw),
        )
        stmt = stmt.where(cond)
        count_stmt = count_stmt.where(cond)
    if subject:
        stmt = stmt.where(QuestionBankItem.subject == subject.strip())
        count_stmt = count_stmt.where(QuestionBankItem.subject == subject.strip())
    if question_no:
        try:
            no = int(question_no.strip())
            stmt = stmt.where(QuestionBankItem.bank_no == no)
            count_stmt = count_stmt.where(QuestionBankItem.bank_no == no)
        except ValueError:
            return {"items": [], "total": 0}

    total_result = await session.execute(count_stmt)
    total = total_result.scalar() or 0
    stmt = stmt.offset((page - 1) * page_size).limit(page_size)
    result = await session.execute(stmt)
    items = result.scalars().all()
    return {"items": [item.to_pydantic() for item in items], "total": total}


@router.post("/question-bank", response_model=QuestionBankItemSchema)
async def add_to_bank(payload: AddToBankRequest, session: AsyncSession = Depends(get_db)):
    # 防御：如果 exam_id 指向不存在的 exam，视为 None
    if payload.exam_id:
        check = await session.execute(select(Exam.id).where(Exam.id == payload.exam_id))
        if not check.scalar_one_or_none():
            payload.exam_id = None
    if payload.exam_id:
        existing = await session.execute(
            select(QuestionBankItem).where(
                QuestionBankItem.exam_id == payload.exam_id,
                QuestionBankItem.question_no == payload.question_no,
            )
        )
    else:
        existing = await session.execute(
            select(QuestionBankItem).where(
                QuestionBankItem.question_text == payload.question_text,
                QuestionBankItem.standard_answer == payload.standard_answer,
            )
        )
    if existing.scalar_one_or_none():
        raise HTTPException(status_code=409, detail="该题目已在题库中")

    # 分配入库序号：最大 bank_no + 1
    max_no_result = await session.execute(select(func.max(QuestionBankItem.bank_no)))
    max_no = max_no_result.scalar() or 0

    item = QuestionBankItem(
        exam_id=payload.exam_id,
        question_no=payload.question_no,
        question_text=payload.question_text,
        standard_answer=payload.standard_answer,
        analysis=payload.analysis,
        exam_filename=payload.exam_filename,
        bank_no=max_no + 1,
        subject=payload.subject or _guess_subject(payload.exam_filename, payload.question_text),
        added_at=datetime.now(),
    )
    session.add(item)
    await session.commit()
    await session.refresh(item)
    return item.to_pydantic()


@router.delete("/question-bank/{item_id}")
async def remove_from_bank(item_id: int, session: AsyncSession = Depends(get_db)):
    item = await session.get(QuestionBankItem, item_id)
    if item:
        await session.delete(item)
        await session.commit()
    return {"status": "ok"}


@router.delete("/question-bank")
async def clear_bank(session: AsyncSession = Depends(get_db)):
    await session.execute(delete(QuestionBankItem))
    await session.commit()
    return {"status": "ok"}


@router.get("/question-bank/all")
async def all_question_bank(session: AsyncSession = Depends(get_db)):
    """返回全部题库题目（不分页，供训练出题使用）。"""
    result = await session.execute(
        select(QuestionBankItem).order_by(QuestionBankItem.bank_no.asc())
    )
    items = result.scalars().all()
    return [item.to_pydantic() for item in items]


@router.post("/question-bank/reclassify")
async def reclassify_question_bank(session: AsyncSession = Depends(get_db)):
    """根据最新关键词库重新分类所有题库题目。"""
    result = await session.execute(select(QuestionBankItem))
    items = result.scalars().all()
    updated = 0
    for item in items:
        new_subj = _guess_subject(question_text=item.question_text)
        if item.subject != new_subj:
            item.subject = new_subj
            updated += 1
    await session.commit()
    return {"status": "ok", "total": len(items), "updated": updated}


def _mock_questions(payload: AIGenerateRequest) -> list[dict]:
    """根据请求参数动态生成 mock 题目（AI不可用时的兜底）。"""
    is_primary = payload.grade == "小学"
    req = (payload.requirement or "").strip()

    if payload.subject == "语文":
        if is_primary:
            templates = [
                {"q": "看拼音写词语：táo huā（    ）", "a": "桃花", "r": "考查拼音拼读和汉字书写，注意'桃'的右边是'兆'。"},
                {"q": "下列词语中加点字读音完全正确的一项是？A. 模范(mú) B. 模样(mó) C. 模范(mó)", "a": "C", "r": "'模范'的'模'读mó，'模样'的'模'读mú，注意多音字辨析。"},
                {"q": "默写李白的《静夜思》。", "a": "床前明月光，疑是地上霜。举头望明月，低头思故乡。", "r": "考查课标必背古诗，注意'疑''霜'二字的书写。"},
                {"q": "把下列'把'字句改成'被'字句：小猫把花瓶打碎了。", "a": "花瓶被小猫打碎了。", "r": "把字句与被字句转换，交换主语和宾语的位置。"},
                {"q": "'他跑得很快'这句话用了什么修辞手法？A. 比喻 B. 拟人 C. 夸张", "a": "C", "r": "考查修辞手法判断，'很快'是夸张表达。"},
                {"q": "选词填空：安静 宁静 平静\n教室里非常（    ），同学们都在认真写作业。", "a": "安静", "r": "近义词辨析：安静强调没有声音，宁静形容环境，平静形容心情。"},
                {"q": "请写出三个描写春天的成语。", "a": "春暖花开、春光明媚、鸟语花香（答案不唯一）", "r": "考查成语积累，注意与主题的关联性。"},
                {"q": "修改病句：经过老师的帮助，我的成绩有了明显的提高和进步。", "a": "经过老师的帮助，我的成绩有了明显的提高。", "r": "语义重复，'提高'和'进步'意思相近，删去其一。"},
                {"q": "阅读短文回答问题（略），请概括这篇文章的主要内容。", "a": "概括要抓住时间、地点、人物、事件四要素。", "r": "考查阅读理解与概括能力，注意用简洁的语言归纳。"},
                {"q": "请用'虽然……但是……'写一句话。", "a": "答案不唯一，如：虽然下雨了，但是我还是坚持去上学。", "r": "考查转折关系的关联词运用，注意前后语义相反。"},
            ]
        else:
            templates = [
                {"q": "下列词语中加点字读音完全正确的一项是？", "a": "B", "r": "A项'殷红'应为yān，C项'拮据'应为jū，D项'栈桥'应为zhàn。"},
                {"q": "默写杜甫《春望》中描写战乱景象的两句。", "a": "国破山河在，城春草木深。", "r": "考查古诗文默写，注意'破''深'二字。"},
                {"q": "'温故而知新'中'故'的意思是？", "a": "旧的知识", "r": "语出《论语》，'故'指学过的知识。"},
                {"q": "下列句子没有语病的一项是？", "a": "通过这次活动，我开阔了眼界。", "r": "考查病句辨析，其余选项均有搭配不当。"},
                {"q": "请写出《陋室铭》的主旨句。", "a": "斯是陋室，惟吾德馨。", "r": "全文主旨，体现作者安贫乐道。"},
                {"q": "'吹面不寒杨柳风'出自哪位诗人之手？", "a": "志南（释志南）", "r": "南宋诗僧志南的《绝句》。"},
                {"q": "下列修辞手法判断错误的是？", "a": "C（比喻应为拟人）", "r": "考查比喻、拟人、排比、夸张的区分。"},
                {"q": "'不以物喜，不以己悲'表现了怎样的精神境界？", "a": "豁达胸怀，不因外物的好坏和个人的得失而或喜或悲。", "r": "出自《岳阳楼记》，考查思想感情分析。"},
                {"q": "下列成语使用恰当的一项是？", "a": "A", "r": "考查成语在具体语境中的运用。注意感情色彩。"},
                {"q": "请翻译：'知之者不如好之者，好之者不如乐之者。'", "a": "知道它的人不如喜爱它的人，喜爱它的人不如以它为乐的人。", "r": "考查文言句子翻译，注意'之'的指代。"},
            ]
    else:
        if is_primary:
            templates = [
                {"q": "计算：125 × 8 ÷ 4 = ?", "a": "250", "r": "先乘后除，125×8=1000，1000÷4=250。"},
                {"q": "把 3/4 和 5/6 通分后比较大小。", "a": "3/4 = 9/12，5/6 = 10/12，所以 3/4 < 5/6", "r": "找分母的最小公倍数12，然后比较分子。"},
                {"q": "一个长方形的长是12cm，宽是8cm，求它的周长和面积。", "a": "周长=40cm，面积=96cm²", "r": "周长=(长+宽)×2=40，面积=长×宽=96。"},
                {"q": "小明买了3支笔和2个本子，笔每支2元，本子每个5元，一共花了多少钱？", "a": "16元", "r": "3×2+2×5=6+10=16元。"},
                {"q": "36和48的最大公因数是多少？", "a": "12", "r": "36=2²×3², 48=2⁴×3，公因数取最小指数：2²×3=12。"},
                {"q": "2.5千克 = （    ）克", "a": "2500", "r": "1千克=1000克，2.5×1000=2500克。"},
                {"q": "解方程：3x + 5 = 20", "a": "x = 5", "r": "移项：3x=15，两边除以3得x=5。"},
                {"q": "三角形三个内角的和是多少度？", "a": "180°", "r": "三角形内角和定理，任意三角形内角和都是180°。"},
                {"q": "一个正方体的棱长是4cm，求它的体积。", "a": "64cm³", "r": "正方体体积=棱长³=4³=64。"},
                {"q": "根据统计图，某班男生25人女生20人，男生比女生多百分之几？", "a": "25%", "r": "(25-20)÷20×100%=25%。"},
            ]
        else:
            templates = [
                {"q": "解方程：3x - 7 = 2x + 5", "a": "x = 12", "r": "移项合并同类项即可。"},
                {"q": "若方程 2x² - 5x + k = 0 有两个相等实数根，求 k。", "a": "k = 25/8", "r": "判别式 Δ = b²-4ac = 25-8k = 0。"},
                {"q": "抛物线 y = x² - 4x + 3 的顶点坐标是？", "a": "(2, -1)", "r": "配方：y = (x-2)² - 1，顶点(2,-1)。"},
                {"q": "在直角三角形中，∠C=90°，AC=3, BC=4，求 AB。", "a": "5", "r": "勾股定理：AB = √(3²+4²) = 5。"},
                {"q": "计算：(-2)³ + √16 - |3-7|", "a": "-3", "r": "=-8+4-4=-8，注意运算顺序。"},
                {"q": "因式分解：x² - 5x + 6", "a": "(x-2)(x-3)", "r": "十字相乘法，找到-2和-3。"},
                {"q": "函数 y = 2x + 1 经过第几象限？", "a": "一、二、三象限", "r": "k=2>0,b=1>0, 过一二三象限。"},
                {"q": "解不等式组：{2x-3>1, x+2≤7}", "a": "2 < x ≤ 5", "r": "分别求解取交集。"},
                {"q": "已知 a² - b² = 21，a + b = 7，求 a - b。", "a": "3", "r": "a²-b²=(a+b)(a-b)，即21=7×(a-b)。"},
                {"q": "计算概率：掷两个骰子，点数和为7的概率。", "a": "1/6", "r": "6/36=1/6，共有6种组合。注意有序性。"},
            ]

    # 如果指定了知识点，给 mock 题目打上标注
    if req:
        prefix = f"【{req}】"
        for t in templates:
            if prefix not in t["q"]:
                t["q"] = f"{prefix} {t['q']}"

    result = []
    for i in range(min(payload.count, len(templates))):
        t = templates[i]
        result.append({
            "question_no": str(i + 1),
            "question_text": t["q"],
            "standard_answer": t["a"],
            "analysis": t["r"],
            "difficulty": payload.difficulty,
        })
    return result


# =====================================================================
# AI 出题接口
# =====================================================================

@router.post("/ai-generate", response_model=AIGenerateResponse)
async def ai_generate_questions(payload: AIGenerateRequest):
    if AGENT_AVAILABLE and homework_agent is not None:
        try:
            questions_data = await homework_agent.generate_questions(
                subject=payload.subject,
                grade=payload.grade,
                question_type=payload.question_type,
                difficulty=payload.difficulty,
                count=payload.count,
                requirement=payload.requirement,
            )
        except Exception as e:
            logger.error(f"AI出题失败: {e}", exc_info=True)
            questions_data = _mock_questions(payload)
    else:
        questions_data = _mock_questions(payload)

    return AIGenerateResponse(
        questions=[AIGeneratedQuestion(**q) for q in questions_data]
    )


@router.get("/health")
async def health_check(session: AsyncSession = Depends(get_db)):
    try:
        from sqlalchemy import text
        await session.execute(text("SELECT 1"))
        db_status = "ok"
    except Exception:
        db_status = "error"
    return {"status": "ok", "db": db_status, "ai_available": AGENT_AVAILABLE, "timestamp": datetime.now().isoformat()}
